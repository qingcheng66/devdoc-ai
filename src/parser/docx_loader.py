from pathlib import Path
from typing import List

from langchain_core.documents import Document
from src.parser.sanitizer import sanitize_text
from src.utils.logger import logger


class DocxLoader:
    def can_load(self, file_path: str) -> bool:
        return Path(file_path).suffix.lower() == ".docx"

    def load(self, file_path: str) -> List[Document]:
        from docx import Document as DocxDocument
        doc = DocxDocument(file_path)
        parts = []

        # Read regular paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                parts.append(text)

        # Read table cells — many docx reports put content in tables
        for table in doc.tables:
            for row in table.rows:
                row_texts = []
                for cell in row.cells:
                    text = cell.text.strip()
                    if text:
                        row_texts.append(text)
                if row_texts:
                    parts.append(" | ".join(row_texts))

        full_text = "\n".join(parts)
        full_text = sanitize_text(full_text)
        if not full_text.strip():
            logger.warning(f"Empty content from {file_path}")
            return []
        logger.info(f"Loaded docx: {Path(file_path).name} ({len(full_text)} chars, {len(parts)} text segments)")
        return [Document(
            page_content=full_text,
            metadata={"source": file_path, "file_type": "docx",
                       "segment_count": len(parts)}
        )]
